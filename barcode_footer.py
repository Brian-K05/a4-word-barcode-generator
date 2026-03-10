"""
A4 Word document generator with barcode in footer.
- Barcode type: Code 128
- Flow: input barcode value → generate Word file with barcode in footer.

Footer layout:
  - Size: H 0.22", W 1.5" (aspect ratio not locked)
  - Position: Horizontal 6.62", Vertical 11.39" (relative to page)
  - Text wrapping: In front of text
  - Page: A4
"""

import io
import os
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml

# EMUs per inch (Word OOXML)
EMU_PER_INCH = 914400

# Layout constants from your spec
BARCODE_WIDTH_IN = 1.5
BARCODE_HEIGHT_IN = 0.22
POS_HORIZONTAL_IN = 6.62
POS_VERTICAL_IN = 11.39


def _emu(inches):
    """Convert inches to EMU (English Metric Units)."""
    return int(inches * EMU_PER_INCH)


def _anchor_xml(pos_x_emu, pos_y_emu, cx_emu, cy_emu, rId):
    """XML for wp:anchor with position relative to page, in front of text (wrapNone, behindDoc=0)."""
    return (
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" '
        'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">\n'
        '  <wp:simplePos x="0" y="0"/>\n'
        '  <wp:positionH relativeFrom="page">\n'
        '    <wp:posOffset>%d</wp:posOffset>\n'
        '  </wp:positionH>\n'
        '  <wp:positionV relativeFrom="page">\n'
        '    <wp:posOffset>%d</wp:posOffset>\n'
        '  </wp:positionV>\n'
        '  <wp:extent cx="%d" cy="%d"/>\n'
        '  <wp:effectExtent l="0" t="0" r="0" b="0"/>\n'
        '  <wp:wrapNone/>\n'
        '  <wp:docPr id="1" name="Barcode"/>\n'
        '  <wp:cNvGraphicFramePr>\n'
        '    <a:graphicFrameLocks noChangeAspect="0" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>\n'
        '  </wp:cNvGraphicFramePr>\n'
        '  <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
        '    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">\n'
        '      <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">\n'
        '        <pic:nvPicPr>\n'
        '          <pic:cNvPr id="0" name="Barcode"/>\n'
        '          <pic:cNvPicPr/>\n'
        '        </pic:nvPicPr>\n'
        '        <pic:blipFill>\n'
        '          <a:blip r:embed="%s"/>\n'
        '          <a:stretch>\n'
        '            <a:fillRect/>\n'
        '          </a:stretch>\n'
        '        </pic:blipFill>\n'
        '        <pic:spPr>\n'
        '          <a:xfrm>\n'
        '            <a:off x="0" y="0"/>\n'
        '            <a:ext cx="%d" cy="%d"/>\n'
        '          </a:xfrm>\n'
        '          <a:prstGeom prst="rect">\n'
        '            <a:avLst/>\n'
        '          </a:prstGeom>\n'
        '        </pic:spPr>\n'
        '      </pic:pic>\n'
        '    </a:graphicData>\n'
        '  </a:graphic>\n'
        '</wp:anchor>'
        % (pos_x_emu, pos_y_emu, cx_emu, cy_emu, rId, cx_emu, cy_emu)
    )


# Barcode image options used for both Word and standalone image (must match)
BARCODE_OPTIONS = {'write_text': False}


def generate_barcode_image_bytes(barcode_data):
    """
    Generate Code 128 barcode image as PNG bytes.
    Use this so the same image can be embedded in the Word doc and saved as a file.
    """
    try:
        import barcode
        from barcode.writer import ImageWriter
    except ImportError:
        raise ImportError(
            "Install python-barcode and Pillow: pip install python-barcode Pillow"
        )
    code = barcode.get('code128', str(barcode_data), writer=ImageWriter())
    buffer = io.BytesIO()
    code.write(buffer, options=BARCODE_OPTIONS)
    buffer.seek(0)
    return buffer.getvalue()


def create_document_with_barcode(barcode_data, output_path, barcode_image_path=None):
    """
    Create an A4 Word document with a barcode image in the footer.

    :param barcode_data: Text to encode in the barcode (e.g. "123456789").
    :param output_path: Path for the output .docx file or file-like object (e.g. BytesIO).
    :param barcode_image_path: Optional path to an existing barcode image file (same as standalone image).
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    if barcode_image_path and os.path.isfile(barcode_image_path):
        with open(barcode_image_path, 'rb') as f:
            image_stream = io.BytesIO(f.read())
    else:
        image_stream = io.BytesIO(generate_barcode_image_bytes(barcode_data))

    footer = section.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    paragraph = footer.paragraphs[0]
    part = paragraph.part

    image_stream.seek(0)
    rId, image = part.get_or_add_image(image_stream)
    cx_emu = _emu(BARCODE_WIDTH_IN)
    cy_emu = _emu(BARCODE_HEIGHT_IN)
    pos_x = _emu(POS_HORIZONTAL_IN)
    pos_y = _emu(POS_VERTICAL_IN)

    anchor_xml_str = _anchor_xml(pos_x, pos_y, cx_emu, cy_emu, rId)
    anchor = parse_xml(anchor_xml_str)

    run = paragraph.add_run()
    run._r.add_drawing(anchor)

    doc.save(output_path)
    return output_path


def generate_word_bytes(barcode_data):
    """Create A4 Word with Code 128 barcode in footer, return .docx as bytes (for download)."""
    buffer = io.BytesIO()
    create_document_with_barcode(barcode_data, buffer, barcode_image_path=None)
    buffer.seek(0)
    return buffer.getvalue()


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description='Input a barcode value → generate A4 Word file with Code 128 barcode in footer.'
    )
    parser.add_argument('barcode_value', nargs='?', help='Barcode value to encode (Code 128)')
    parser.add_argument('-o', '--output', default=None, help='Output .docx path (default: barcode_<value>.docx)')
    args = parser.parse_args()

    if not args.barcode_value:
        args.barcode_value = input('Enter barcode value: ').strip()
    if not args.barcode_value:
        parser.error('Barcode value is required.')

    output_path = args.output
    if not output_path:
        safe = "".join(c if c.isalnum() or c in '-_' else '_' for c in args.barcode_value)[:50]
        output_path = f'barcode_{safe}.docx'

    create_document_with_barcode(args.barcode_value, output_path)
    print('Generated:', output_path)


if __name__ == '__main__':
    main()
